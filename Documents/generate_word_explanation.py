from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import RGBColor

# Create a new Word document
doc = Document()

# Title
doc.add_heading('AI Chatbot MVP – Beginner\'s Guide', 0)

doc.add_paragraph('This document explains the AI Chatbot MVP project in simple terms, with diagrams and screenshot placeholders. It is designed for absolute beginners.')

def add_section(title):
    doc.add_page_break()
    doc.add_heading(title, level=1)

def add_diagram(diagram, caption=None):
    p = doc.add_paragraph(diagram)
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    if caption:
        doc.add_paragraph(caption).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

def add_bullet_list(items):
    for item in items:
        doc.add_paragraph(item, style='List Bullet')

def add_callout(text, color=RGBColor(0,0,255)):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.color.rgb = color
    p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

# Table of Contents
add_section('Table of Contents')
doc.add_paragraph('1. Project Overview')
doc.add_paragraph('2. How the Chatbot Works (with Diagram)')
doc.add_paragraph('3. Project Structure & Files (with Diagram)')
doc.add_paragraph('4. Backend (Server) Explanation')
doc.add_paragraph('5. Frontend (User Interface) Explanation (with UI Diagram)')
doc.add_paragraph('6. How Everything Connects (with Flowchart)')
doc.add_paragraph('7. How to Run the Project')
doc.add_paragraph('8. Screenshots (Replace Placeholders)')

# 1. Project Overview
add_section('1. Project Overview')
doc.add_paragraph(
    'This project is a simple web-based chatbot powered by artificial intelligence (AI). '
    'It allows users to type questions or messages and get smart replies from an AI assistant. '
    'The project is designed to be easy to understand and run, even for beginners.'
)

# 2. How the Chatbot Works
add_section('2. How the Chatbot Works')
doc.add_paragraph('Step-by-step Flow:')
add_bullet_list([
    "The user types a message in the chat window on a web page.",
    "The message is sent to a server (the backend) running on your computer.",
    "The server asks the AI (from OpenAI) for a reply.",
    "The server sends the AI's reply back to the web page.",
    "The reply appears in the chat window for the user to read."
])
add_diagram(
    '+---------+        +-----------+        +--------+\n'
    '|  User   | <----> |  Webpage  | <----> | Server |\n'
    '+---------+        +-----------+        +--------+\n'
    '                                         |\n'
    '                                         v\n'
    '                                    +----------+\n'
    '                                    |   OpenAI |\n'
    '                                    +----------+',
    caption='Diagram: Chatbot Flow'
)

# 3. Project Structure & Files
add_section('3. Project Structure & Files')
doc.add_paragraph('Folder Structure:')
doc.add_paragraph(
    'SaturnIQ/\n'
    '├── backend/\n'
    '│   └── app.py\n'
    '├── frontend/\n'
    '│   └── index.html\n'
    '└── Documents/\n'
    '    └── AI_Chatbot_MVP_Explanation.docx'
)
doc.add_paragraph('What Each Folder/File Is For:')
add_bullet_list([
    'backend/: Contains the server code (Python)',
    '  - app.py: The main server script that talks to the AI and handles messages.',
    'frontend/: Contains the web page files',
    '  - index.html: The main web page you see in your browser.',
    'Documents/: Contains documentation and explanations',
    '  - AI_Chatbot_MVP_Explanation.docx: This explanation file.'
])

# 4. Backend (Server) Explanation
add_section('4. Backend (Server) Explanation')
doc.add_paragraph('The backend is a Python program that runs on your computer. It uses a tool called Flask to create a simple web server. It listens for messages from the web page (frontend). When it gets a message, it sends it to the AI and waits for a reply. Once the AI replies, the backend sends the answer back to the web page. The backend also handles errors, like if the AI is not available.')
doc.add_paragraph('Main Parts of the Backend:')
add_bullet_list([
    'Flask: Lets Python act as a web server.',
    'CORS: Allows the web page to talk to the backend, even though they run separately.',
    'OpenAI: Connects to the AI service to get smart replies.',
    'API Key: A secret code that lets the backend use the AI service.'
])
doc.add_paragraph('What Happens When You Send a Message:')
add_bullet_list([
    "The backend receives your message from the web page.",
    "It sends your message to the AI.",
    "It waits for the AI to reply.",
    "It sends the AI's reply back to the web page.",
    "If something goes wrong, it sends an error message instead."
])

# 5. Frontend (User Interface) Explanation
add_section('5. Frontend (User Interface) Explanation')
doc.add_paragraph('The frontend is a web page you open in your browser. It shows a chat window where you can type messages and see replies. It uses HTML for structure, CSS for styling, and JavaScript for making things interactive.')
doc.add_paragraph('Main Parts of the Frontend:')
add_bullet_list([
    'Chat Window: Shows the conversation between you and the AI.',
    'Input Box: Where you type your message.',
    'Send Button: Click to send your message.',
    'Clear Chat Button: Clears the chat history.',
    'Theme Toggle: Switches between light and dark mode.',
    'Loading Spinner: Shows when waiting for the AI to reply.',
    'Copy Button: Lets you copy a message to your clipboard.',
    'Toast Notification: Shows small pop-up messages for errors or actions.'
])
add_diagram(
    '+-----------------------------------+\n'
    '|           Chat Window             |\n'
    '|  [User/AI messages appear here]   |\n'
    '+-----------------------------------+\n'
    '| [Input Box] [Send] [Clear] [🌙]   |\n'
    '+-----------------------------------+',
    caption='UI Layout Diagram'
)

# 6. How Everything Connects
add_section('6. How Everything Connects')
doc.add_paragraph('Flowchart:')
add_diagram(
    '[User] --types--> [Webpage] --sends message--> [Backend]\n'
    '    ^                                         |\n'
    '    |                                         v\n'
    '[AI reply] <--from AI-- [Backend] <-- [OpenAI API]'
)
doc.add_paragraph('The frontend (web page) and backend (server) talk to each other using the internet, but both run on your computer. The frontend sends messages to the backend using a special address (http://localhost:5000/api/chat). The backend talks to the AI and sends the reply back to the frontend. The frontend updates the chat window with the new message.')

# 7. How to Run the Project
add_section('7. How to Run the Project')
add_bullet_list([
    'Start the backend (server): Open a terminal and run the backend Python script.',
    'Start the frontend (web page): Open another terminal and run the command to start the web server for the frontend.',
    'Open your browser and go to http://localhost:8000 to use the chatbot.'
])

# 8. Screenshots (Replace Placeholders)
add_section('8. Screenshots (Replace Placeholders)')
doc.add_paragraph('Tip: To take a screenshot, press PrtScn (Print Screen) or use the Snipping Tool on Windows. Paste (Ctrl+V) or insert the image into the document where indicated below.')

def screenshot_placeholder(title):
    doc.add_paragraph(title, style='Heading 2')
    doc.add_paragraph('[Insert screenshot here]', style='Intense Quote')

screenshot_placeholder('a) Chat Window (Initial State)')
screenshot_placeholder('b) Sending a Message')
screenshot_placeholder('c) AI Reply')
screenshot_placeholder('d) Error State')
screenshot_placeholder('e) Dark Mode')

# Callouts
add_section('Callouts')
add_callout('Note: The diagrams above are for illustration. You can draw them in Word using shapes for a neater look.', color=RGBColor(0,0,255))
add_callout('Warning: Never share your API key publicly.', color=RGBColor(255,0,0))
add_callout('Tip: Replace all screenshot placeholders with your own screenshots for a complete guide.', color=RGBColor(0,128,0))

# Save the document
output_path = 'Documents/AI_Chatbot_MVP_Explanation.docx'
doc.save(output_path)
print(f'Word document generated: {output_path}') 